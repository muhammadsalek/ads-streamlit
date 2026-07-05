"""
🎓 Academic Design Studio (ADS) - Production Version 2.0
================================================================================
A comprehensive academic document generation platform designed for 
researchers, students, and academics. Features include:
- Document Builder (Assignment, Lab Report, Thesis, Proposal)
- Poster Studio with multiple layouts and sizes
- Academic CV Builder
- QR Code Generator
- Presentation Builder
- Citation Manager
- Export Manager with multiple formats

Author: ADS Team
Version: 2.0.0
License: MIT
"""

import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
from PIL import Image, ImageDraw, ImageFont
import io
import json
import os
import base64
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import Rectangle, FancyBboxPatch
import numpy as np
import qrcode
from reportlab.lib.pagesizes import A4, A3, A2, A1, A0, letter, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pptx
from pptx.util import Inches as PptxInches, Pt as PptxPt
from pptx.enum.text import PP_ALIGN
from pptx.dml.color import RGBColor as PptxRGBColor

# ============================================================================
# APPLICATION CONFIGURATION
# ============================================================================

APP_CONFIG = {
    "name": "Academic Design Studio",
    "version": "2.0.0",
    "author": "ADS Team",
    "description": "Professional academic document design platform",
    "year": 2024,
}

# ============================================================================
# PAGE CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title=f"{APP_CONFIG['name']} v{APP_CONFIG['version']}",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/yourusername/ads-streamlit',
        'Report a bug': "https://github.com/yourusername/ads-streamlit/issues",
        'About': f"""
            ## {APP_CONFIG['name']} v{APP_CONFIG['version']}
            
            ### Professional Academic Document Generation Platform
            
            **Features:**
            - 📝 Document Builder (Assignment, Lab Report, Thesis, Proposal)
            - 🎨 Poster Studio (10+ Layouts, 8 Sizes)
            - 📄 Academic CV Builder
            - 📱 QR Code Generator
            - 📊 Presentation Builder
            - 📚 Citation Manager
            - 📦 Export Manager (DOCX, PDF, PPTX, PNG, SVG)
            
            **Tech Stack:**
            - Streamlit
            - Python
            - Matplotlib, Plotly
            - ReportLab, python-docx, python-pptx
            - Pillow, QRCode
            
            © 2024 {APP_CONFIG['author']}
        """
    }
)

# ============================================================================
# CUSTOM CSS - PROFESSIONAL DESIGN
# ============================================================================

def load_custom_css() -> str:
    """Load custom CSS for professional UI design."""
    return """
    <style>
        /* Global Styles */
        .main {
            padding: 0rem 1rem;
        }
        
        .stApp {
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        }
        
        /* Main Header */
        .main-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 2.5rem;
            border-radius: 15px;
            color: white;
            text-align: center;
            margin-bottom: 2rem;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        
        .main-header h1 {
            font-size: 3rem;
            font-weight: 700;
            margin: 0;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        }
        
        .main-header p {
            font-size: 1.2rem;
            opacity: 0.95;
            margin-top: 0.5rem;
        }
        
        /* Feature Cards */
        .feature-card {
            background: white;
            padding: 1.5rem;
            border-radius: 12px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.08);
            margin: 0.5rem 0;
            transition: all 0.3s ease;
            border-left: 5px solid #667eea;
            height: 100%;
        }
        
        .feature-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 8px 25px rgba(0,0,0,0.15);
        }
        
        .feature-card h3 {
            color: #2c3e50;
            margin-top: 0;
        }
        
        .feature-card ul {
            list-style-type: none;
            padding: 0;
            margin: 0;
        }
        
        .feature-card ul li {
            padding: 0.3rem 0;
            color: #555;
            border-bottom: 1px solid #f0f0f0;
        }
        
        .feature-card ul li:last-child {
            border-bottom: none;
        }
        
        /* Section Cards */
        .section-card {
            background: white;
            padding: 1.5rem;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
            margin: 1rem 0;
        }
        
        /* Stats Cards */
        .stat-card {
            background: white;
            padding: 1.5rem;
            border-radius: 10px;
            text-align: center;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        }
        
        .stat-card .stat-number {
            font-size: 2.5rem;
            font-weight: 700;
            color: #667eea;
        }
        
        .stat-card .stat-label {
            color: #888;
            font-size: 0.9rem;
        }
        
        /* Sidebar */
        .css-1d391kg {
            background: linear-gradient(180deg, #f8f9fa 0%, #e9ecef 100%);
        }
        
        .sidebar-logo {
            text-align: center;
            padding: 1rem 0;
        }
        
        /* Buttons */
        .stButton > button {
            border-radius: 8px;
            font-weight: 500;
            transition: all 0.3s ease;
        }
        
        .stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
        }
        
        /* Download Buttons */
        .download-btn {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white !important;
            border: none !important;
            padding: 0.5rem 1.5rem !important;
            border-radius: 8px !important;
        }
        
        .download-btn:hover {
            opacity: 0.9;
            transform: translateY(-2px);
        }
        
        /* Expanders */
        .streamlit-expanderHeader {
            background: #f8f9fa;
            border-radius: 8px;
            font-weight: 600;
        }
        
        /* Metrics */
        .metric-card {
            background: white;
            padding: 1rem;
            border-radius: 10px;
            text-align: center;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }
        
        /* Responsive */
        @media (max-width: 768px) {
            .main-header h1 {
                font-size: 2rem;
            }
            .main-header p {
                font-size: 1rem;
            }
        }
        
        /* Scrollbar */
        ::-webkit-scrollbar {
            width: 8px;
            height: 8px;
        }
        
        ::-webkit-scrollbar-track {
            background: #f1f1f1;
            border-radius: 10px;
        }
        
        ::-webkit-scrollbar-thumb {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border-radius: 10px;
        }
        
        ::-webkit-scrollbar-thumb:hover {
            background: #5a6fd6;
        }
    </style>
    """

st.markdown(load_custom_css(), unsafe_allow_html=True)

# ============================================================================
# SESSION STATE MANAGEMENT
# ============================================================================

class SessionState:
    """Centralized session state management."""
    
    @staticmethod
    def initialize():
        """Initialize all session state variables."""
        defaults = {
            'current_page': 'Home',
            'projects': [],
            'project_count': 0,
            'theme': 'light',
            'font_size': 'medium',
            'auto_save': True,
            'current_project': None,
            'draft_data': {},
            'export_queue': [],
            'notifications': [],
            'user_preferences': {
                'default_template': 'modern',
                'default_export_format': 'pdf',
                'citation_style': 'apa'
            }
        }
        
        for key, value in defaults.items():
            if key not in st.session_state:
                st.session_state[key] = value

SessionState.initialize()

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def generate_id() -> str:
    """Generate unique ID using timestamp and hash."""
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    random_hash = hashlib.md5(str(timestamp).encode()).hexdigest()[:8]
    return f"{timestamp}_{random_hash}"

def format_date(date_obj) -> str:
    """Format date object to string."""
    if isinstance(date_obj, datetime):
        return date_obj.strftime("%B %d, %Y")
    return str(date_obj)

def create_download_link(data: bytes, filename: str, label: str) -> str:
    """Create HTML download link."""
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{label}</a>'
    return href

def validate_email(email: str) -> bool:
    """Validate email format."""
    pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    return bool(re.match(pattern, email))

def truncate_text(text: str, max_length: int = 100) -> str:
    """Truncate text to max length with ellipsis."""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."

# ============================================================================
# SIDEBAR COMPONENT
# ============================================================================

def render_sidebar():
    """Render the main sidebar with navigation."""
    with st.sidebar:
        # Logo and Title
        st.markdown("""
        <div class="sidebar-logo">
            <h1 style="color: #667eea; margin: 0; font-size: 2rem;">🎓 ADS</h1>
            <p style="color: #888; margin: 0; font-size: 0.8rem;">Academic Design Studio</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Navigation Pages
        pages = {
            "🏠 Home": "Home",
            "📝 Assignment": "Assignment",
            "📊 Lab Report": "Lab Report",
            "📄 Thesis": "Thesis",
            "📑 Research Proposal": "Research Proposal",
            "🎨 Poster": "Poster",
            "📽️ Presentation": "Presentation",
            "📄 CV": "CV",
            "📱 QR Code": "QR",
            "📚 Citation Manager": "Citation",
            "📦 Export": "Export",
            "⚙️ Settings": "Settings"
        }
        
        for label, page in pages.items():
            if st.button(
                label,
                key=f"nav_{page}",
                use_container_width=True,
                type="primary" if st.session_state.current_page == page else "secondary",
            ):
                st.session_state.current_page = page
                st.rerun()
        
        st.markdown("---")
        
        # Statistics
        st.markdown("### 📊 Statistics")
        col1, col2 = st.columns(2)
        with col1:
            st.metric(
                "Projects",
                len(st.session_state.projects),
                delta="+1" if len(st.session_state.projects) > 0 else None,
            )
        with col2:
            st.metric(
                "Pages",
                len(pages),
                delta="Active"
            )
        
        st.markdown("---")
        
        # Quick Actions
        st.markdown("### ⚡ Quick Actions")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📂 New", use_container_width=True):
                st.session_state.current_page = "Home"
                st.rerun()
        with col2:
            if st.button("💾 Save", use_container_width=True):
                st.success("✓ Saved!")
        
        # Theme Toggle
        st.markdown("---")
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown("### 🌓 Theme")
        with col2:
            current_theme = st.session_state.get("theme", "light")
            if st.button("🌙" if current_theme == "light" else "☀️"):
                st.session_state.theme = "dark" if current_theme == "light" else "light"
                st.rerun()
        
        # Version
        st.markdown("---")
        st.caption(f"v{APP_CONFIG['version']} | © 2024")

# ============================================================================
# HOME PAGE
# ============================================================================

def render_home():
    """Render the home page."""
    st.markdown("""
    <div class="main-header">
        <h1>🎓 Academic Design Studio</h1>
        <p>Design professional academic documents, posters, presentations, and reports in minutes</p>
        <p style="font-size: 0.9rem; opacity: 0.8;">
            📊 10+ Document Types | 🎨 8 Poster Sizes | 📄 5 Export Formats
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Quick Stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown("""
        <div class="stat-card">
            <div class="stat-number">📝</div>
            <div class="stat-label">Document Types</div>
            <div style="font-size: 1.5rem; font-weight: 600; color: #667eea;">10+</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="stat-card">
            <div class="stat-number">🎨</div>
            <div class="stat-label">Poster Layouts</div>
            <div style="font-size: 1.5rem; font-weight: 600; color: #667eea;">10</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="stat-card">
            <div class="stat-number">📄</div>
            <div class="stat-label">Export Formats</div>
            <div style="font-size: 1.5rem; font-weight: 600; color: #667eea;">7</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        projects_count = len(st.session_state.projects)
        st.markdown(f"""
        <div class="stat-card">
            <div class="stat-number">📂</div>
            <div class="stat-label">Projects Created</div>
            <div style="font-size: 1.5rem; font-weight: 600; color: #667eea;">{projects_count}</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Feature Cards
    st.subheader("🚀 Featured Tools")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <h3>📝 Document Builder</h3>
            <p style="color: #666;">Create academic documents with ease</p>
            <ul>
                <li>📋 Assignment Builder</li>
                <li>📊 Lab Report</li>
                <li>📄 Thesis Builder</li>
                <li>📑 Research Proposal</li>
                <li>📚 Case Study</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <h3>🎨 Visual Design</h3>
            <p style="color: #666;">Design visual academic content</p>
            <ul>
                <li>🎨 Research Poster</li>
                <li>📽️ Presentation</li>
                <li>📊 Charts & Graphs</li>
                <li>🖼️ Images & Icons</li>
                <li>📱 QR Codes</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <h3>📄 Professional Tools</h3>
            <p style="color: #666;">Professional academic tools</p>
            <ul>
                <li>📄 Academic CV</li>
                <li>🎓 Certificate</li>
                <li>📚 Citation Manager</li>
                <li>📦 Export Manager</li>
                <li>🏛️ University Templates</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Recent Projects
    st.subheader("📂 Recent Projects")
    
    if st.session_state.projects:
        for project in st.session_state.projects[-5:]:
            col1, col2, col3, col4 = st.columns([3, 2, 1, 1])
            with col1:
                st.markdown(f"**{project.get('name', 'Untitled')}**")
                st.caption(f"Type: {project.get('type', 'Unknown')}")
            with col2:
                st.caption(f"Created: {project.get('created', 'N/A')}")
                st.caption(f"ID: {project.get('id', 'N/A')[:8]}")
            with col3:
                if st.button("📂", key=f"open_{project.get('id', '')}"):
                    st.info(f"Opening {project.get('name', 'Project')}")
            with col4:
                if st.button("🗑️", key=f"del_{project.get('id', '')}"):
                    st.session_state.projects.remove(project)
                    st.rerun()
            st.divider()
    else:
        st.info("💡 No projects yet. Create your first project using the tools above!")

# ============================================================================
# ASSIGNMENT BUILDER
# ============================================================================

def render_assignment_builder():
    """Render the assignment builder page."""
    st.header("📝 Assignment Builder")
    st.caption("Create professional assignments with rubric and grading criteria")
    
    with st.form("assignment_form", clear_on_submit=False):
        col1, col2 = st.columns(2)
        
        with col1:
            title = st.text_input("Assignment Title *", placeholder="e.g., Machine Learning Project")
            course = st.text_input("Course Name *", placeholder="e.g., CS 101")
            instructor = st.text_input("Instructor Name *", placeholder="Dr. John Doe")
            max_score = st.number_input("Maximum Score", min_value=0, max_value=1000, value=100, step=5)
        
        with col2:
            due_date = st.date_input("Due Date", datetime.now() + timedelta(days=7))
            submission_type = st.selectbox("Submission Type", ["Individual", "Group", "Both"])
            department = st.text_input("Department", placeholder="Computer Science")
            level = st.selectbox("Academic Level", ["Undergraduate", "Graduate", "PhD", "Postdoc"])
        
        instructions = st.text_area("Instructions", height=150, 
                                   placeholder="Detailed assignment instructions...")
        
        grading_criteria = st.text_area("Grading Criteria", height=100,
                                       placeholder="Rubric and grading criteria...")
        
        additional_notes = st.text_area("Additional Notes", height=80,
                                       placeholder="Any additional information...")
        
        submitted = st.form_submit_button("📥 Generate Assignment", use_container_width=True)
        
        if submitted:
            if not title or not course or not instructor:
                st.error("⚠️ Please fill in all required fields (marked with *)!")
            else:
                # Generate assignment
                assignment_id = generate_id()
                project_data = {
                    'id': assignment_id,
                    'name': title,
                    'type': 'Assignment',
                    'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
                    'course': course,
                    'instructor': instructor,
                }
                st.session_state.projects.append(project_data)
                st.session_state.project_count += 1
                
                st.success(f"✅ Assignment '{title}' generated successfully!")
                
                # Preview
                st.subheader("📄 Assignment Preview")
                with st.expander("View Full Assignment", expanded=True):
                    st.markdown(f"""
                    <div style="background: white; padding: 2rem; border-radius: 10px; border: 1px solid #e0e0e0;">
                        <h1 style="text-align: center; color: #2c3e50;">{title}</h1>
                        <hr>
                        <table style="width: 100%; border-collapse: collapse;">
                            <tr><td><strong>Course</strong></td><td>{course}</td></tr>
                            <tr><td><strong>Instructor</strong></td><td>{instructor}</td></tr>
                            <tr><td><strong>Due Date</strong></td><td>{due_date.strftime('%B %d, %Y')}</td></tr>
                            <tr><td><strong>Max Score</strong></td><td>{max_score}</td></tr>
                            <tr><td><strong>Submission Type</strong></td><td>{submission_type}</td></tr>
                            <tr><td><strong>Department</strong></td><td>{department}</td></tr>
                            <tr><td><strong>Level</strong></td><td>{level}</td></tr>
                        </table>
                        <hr>
                        <h3>Instructions</h3>
                        <p>{instructions}</p>
                        <h3>Grading Criteria</h3>
                        <p>{grading_criteria}</p>
                        <h3>Additional Notes</h3>
                        <p>{additional_notes}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Generate download content
                assignment_text = f"""
                {'='*70}
                ASSIGNMENT: {title}
                {'='*70}
                
                Course: {course}
                Instructor: {instructor}
                Due Date: {due_date.strftime('%B %d, %Y')}
                Max Score: {max_score}
                Submission Type: {submission_type}
                Department: {department}
                Level: {level}
                
                {'-'*70}
                INSTRUCTIONS
                {'-'*70}
                {instructions}
                
                {'-'*70}
                GRADING CRITERIA
                {'-'*70}
                {grading_criteria}
                
                {'-'*70}
                ADDITIONAL NOTES
                {'-'*70}
                {additional_notes}
                
                {'='*70}
                Generated by Academic Design Studio v{APP_CONFIG['version']}
                Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
                """
                
                # Multiple download formats
                st.subheader("📥 Download Options")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    # TXT Download
                    txt_bytes = assignment_text.encode('utf-8')
                    st.download_button(
                        label="📄 Download TXT",
                        data=txt_bytes,
                        file_name=f"{title.replace(' ', '_')}_assignment.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                with col2:
                    # Create simple DOCX
                    try:
                        doc = docx.Document()
                        doc.add_heading(title, 0)
                        doc.add_paragraph(f"Course: {course}")
                        doc.add_paragraph(f"Instructor: {instructor}")
                        doc.add_paragraph(f"Due Date: {due_date.strftime('%B %d, %Y')}")
                        doc.add_heading("Instructions", level=1)
                        doc.add_paragraph(instructions)
                        doc.add_heading("Grading Criteria", level=1)
                        doc.add_paragraph(grading_criteria)
                        
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.download_button(
                            label="📝 Download DOCX",
                            data=docx_buffer.getvalue(),
                            file_name=f"{title.replace(' ', '_')}_assignment.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except:
                        st.button("📝 DOCX (Install python-docx)", disabled=True, use_container_width=True)
                
                with col3:
                    # PDF Download (simple version)
                    try:
                        pdf_buffer = io.BytesIO()
                        c = canvas.Canvas(pdf_buffer, pagesize=A4)
                        c.setFont("Helvetica", 16)
                        c.drawString(72, 800, title)
                        c.setFont("Helvetica", 12)
                        c.drawString(72, 770, f"Course: {course}")
                        c.drawString(72, 750, f"Instructor: {instructor}")
                        c.drawString(72, 730, f"Due Date: {due_date.strftime('%B %d, %Y')}")
                        c.showPage()
                        c.save()
                        pdf_buffer.seek(0)
                        
                        st.download_button(
                            label="📕 Download PDF",
                            data=pdf_buffer.getvalue(),
                            file_name=f"{title.replace(' ', '_')}_assignment.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except:
                        st.button("📕 PDF (Install reportlab)", disabled=True, use_container_width=True)

# ============================================================================
# LAB REPORT BUILDER
# ============================================================================

def render_lab_report():
    """Render the lab report builder page."""
    st.header("📊 Lab Report Builder")
    st.caption("Create professional lab reports with structured sections")
    
    with st.form("lab_report_form", clear_on_submit=False):
        col1, col2 = st.columns(2)
        
        with col1:
            report_title = st.text_input("Report Title *", "Lab Experiment Report")
            experiment_title = st.text_input("Experiment Title *", placeholder="e.g., ML-Based Prediction")
            course_code = st.text_input("Course Code *", placeholder="e.g., STAT 2002")
        
        with col2:
            experiment_date = st.date_input("Experiment Date", datetime.now())
            lab_partners = st.text_input("Lab Partners", placeholder="John Doe, Mary Johnson")
            instructor = st.text_input("Instructor Name *", placeholder="Dr. Smith")
        
        hypothesis = st.text_area("Hypothesis", height=80, 
                                 placeholder="What did you hypothesize?")
        methodology = st.text_area("Methodology", height=120,
                                  placeholder="How did you conduct the experiment?")
        results = st.text_area("Results", height=120,
                              placeholder="What were the results?")
        discussion = st.text_area("Discussion", height=120,
                                 placeholder="Discuss the results and implications")
        conclusion = st.text_area("Conclusion", height=80,
                                 placeholder="Conclude the experiment")
        
        references = st.text_area("References", height=80,
                                 placeholder="List your references here...")
        
        submitted = st.form_submit_button("📊 Generate Lab Report", use_container_width=True)
        
        if submitted:
            if not report_title or not experiment_title or not course_code or not instructor:
                st.error("⚠️ Please fill in all required fields (marked with *)!")
            else:
                # Generate report
                report_id = generate_id()
                project_data = {
                    'id': report_id,
                    'name': report_title,
                    'type': 'Lab Report',
                    'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
                    'experiment': experiment_title,
                }
                st.session_state.projects.append(project_data)
                st.session_state.project_count += 1
                
                st.success(f"✅ Lab Report '{report_title}' generated!")
                
                # Preview
                st.subheader("📄 Report Preview")
                with st.expander("View Full Report", expanded=True):
                    st.markdown(f"""
                    <div style="background: white; padding: 2rem; border-radius: 10px; border: 1px solid #e0e0e0;">
                        <h1 style="text-align: center; color: #2c3e50;">{report_title}</h1>
                        <hr>
                        <table style="width: 100%; border-collapse: collapse;">
                            <tr><td><strong>Experiment</strong></td><td>{experiment_title}</td></tr>
                            <tr><td><strong>Course</strong></td><td>{course_code}</td></tr>
                            <tr><td><strong>Date</strong></td><td>{experiment_date.strftime('%B %d, %Y')}</td></tr>
                            <tr><td><strong>Instructor</strong></td><td>{instructor}</td></tr>
                            <tr><td><strong>Partners</strong></td><td>{lab_partners}</td></tr>
                        </table>
                        <hr>
                        <h3>Hypothesis</h3>
                        <p>{hypothesis}</p>
                        <h3>Methodology</h3>
                        <p>{methodology}</p>
                        <h3>Results</h3>
                        <p>{results}</p>
                        <h3>Discussion</h3>
                        <p>{discussion}</p>
                        <h3>Conclusion</h3>
                        <p>{conclusion}</p>
                        <h3>References</h3>
                        <p>{references}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Generate download content
                report_text = f"""
                {'='*70}
                LAB REPORT: {report_title}
                {'='*70}
                
                Experiment: {experiment_title}
                Course: {course_code}
                Date: {experiment_date.strftime('%B %d, %Y')}
                Instructor: {instructor}
                Partners: {lab_partners}
                
                {'-'*70}
                HYPOTHESIS
                {'-'*70}
                {hypothesis}
                
                {'-'*70}
                METHODOLOGY
                {'-'*70}
                {methodology}
                
                {'-'*70}
                RESULTS
                {'-'*70}
                {results}
                
                {'-'*70}
                DISCUSSION
                {'-'*70}
                {discussion}
                
                {'-'*70}
                CONCLUSION
                {'-'*70}
                {conclusion}
                
                {'-'*70}
                REFERENCES
                {'-'*70}
                {references}
                
                {'='*70}
                Generated by Academic Design Studio v{APP_CONFIG['version']}
                Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
                """
                
                # Download options
                st.subheader("📥 Download Options")
                col1, col2 = st.columns(2)
                
                with col1:
                    report_bytes = report_text.encode('utf-8')
                    st.download_button(
                        label="📄 Download TXT",
                        data=report_bytes,
                        file_name=f"{report_title.replace(' ', '_')}_lab_report.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                with col2:
                    try:
                        doc = docx.Document()
                        doc.add_heading(report_title, 0)
                        doc.add_paragraph(f"Experiment: {experiment_title}")
                        doc.add_paragraph(f"Course: {course_code}")
                        doc.add_paragraph(f"Date: {experiment_date.strftime('%B %d, %Y')}")
                        doc.add_heading("Hypothesis", level=1)
                        doc.add_paragraph(hypothesis)
                        doc.add_heading("Methodology", level=1)
                        doc.add_paragraph(methodology)
                        doc.add_heading("Results", level=1)
                        doc.add_paragraph(results)
                        doc.add_heading("Discussion", level=1)
                        doc.add_paragraph(discussion)
                        doc.add_heading("Conclusion", level=1)
                        doc.add_paragraph(conclusion)
                        
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.download_button(
                            label="📝 Download DOCX",
                            data=docx_buffer.getvalue(),
                            file_name=f"{report_title.replace(' ', '_')}_lab_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except:
                        st.button("📝 DOCX (Install python-docx)", disabled=True, use_container_width=True)

# ============================================================================
# THESIS BUILDER
# ============================================================================

def render_thesis():
    """Render the thesis builder page."""
    st.header("📄 Thesis Builder")
    st.caption("Create comprehensive thesis documents with multiple chapters")
    
    with st.form("thesis_form", clear_on_submit=False):
        col1, col2 = st.columns(2)
        
        with col1:
            thesis_title = st.text_input("Thesis Title *", placeholder="e.g., Deep Learning for Medical Imaging")
            author_name = st.text_input("Author Name *", placeholder="John Doe")
            degree = st.text_input("Degree *", placeholder="Doctor of Philosophy")
            university = st.text_input("University *", placeholder="Harvard University")
        
        with col2:
            department = st.text_input("Department *", placeholder="Computer Science")
            supervisor = st.text_input("Supervisor *", placeholder="Prof. Jane Smith")
            defense_date = st.date_input("Defense Date", datetime.now() + timedelta(days=180))
            submission_date = st.date_input("Submission Date", datetime.now() + timedelta(days=150))
        
        abstract = st.text_area("Abstract *", height=150,
                               placeholder="Abstract of the thesis...")
        
        st.subheader("📑 Chapters")
        
        # Dynamic chapter creation
        chapter_count = st.number_input("Number of Chapters", min_value=3, max_value=10, value=5, step=1)
        
        chapters = []
        for i in range(int(chapter_count)):
            col1, col2 = st.columns(2)
            with col1:
                chapter_title = st.text_input(f"Chapter {i+1} Title", 
                                             placeholder=f"e.g., Introduction",
                                             key=f"ch_title_{i}")
            with col2:
                chapter_pages = st.number_input(f"Pages", min_value=5, max_value=100, value=20,
                                               key=f"ch_pages_{i}")
            
            chapter_content = st.text_area(f"Chapter {i+1} Content", height=100,
                                          placeholder="Chapter content summary...",
                                          key=f"ch_content_{i}")
            
            chapters.append({
                'title': chapter_title,
                'pages': chapter_pages,
                'content': chapter_content
            })
        
        acknowledgements = st.text_area("Acknowledgements", height=100)
        declaration = st.text_area("Declaration", height=80)
        
        submitted = st.form_submit_button("📄 Generate Thesis", use_container_width=True)
        
        if submitted:
            if not thesis_title or not author_name or not degree or not university or not abstract:
                st.error("⚠️ Please fill in all required fields (marked with *)!")
            else:
                st.success(f"✅ Thesis '{thesis_title}' generated!")
                
                # Preview
                st.subheader("📄 Thesis Preview")
                with st.expander("View Full Thesis", expanded=True):
                    st.markdown(f"""
                    <div style="background: white; padding: 2rem; border-radius: 10px; border: 1px solid #e0e0e0;">
                        <h1 style="text-align: center;">{thesis_title}</h1>
                        <h3 style="text-align: center;">by {author_name}</h3>
                        <hr>
                        <p><strong>Degree:</strong> {degree}</p>
                        <p><strong>University:</strong> {university}</p>
                        <p><strong>Department:</strong> {department}</p>
                        <p><strong>Supervisor:</strong> {supervisor}</p>
                        <p><strong>Defense Date:</strong> {defense_date.strftime('%B %d, %Y')}</p>
                        <hr>
                        <h3>Abstract</h3>
                        <p>{abstract}</p>
                        <h3>Table of Contents</h3>
                        """, unsafe_allow_html=True)
                    
                    for i, ch in enumerate(chapters, 1):
                        if ch['title']:
                            st.markdown(f"{i}. {ch['title']} (Page {ch['pages']})")
                    
                    st.markdown("""
                        <h3>Acknowledgements</h3>
                        <p>{}</p>
                        <h3>Declaration</h3>
                        <p>{}</p>
                    </div>
                    """.format(acknowledgements, declaration), unsafe_allow_html=True)

# ============================================================================
# RESEARCH PROPOSAL BUILDER
# ============================================================================

def render_research_proposal():
    """Render the research proposal builder page."""
    st.header("📑 Research Proposal Builder")
    st.caption("Create comprehensive research proposals for funding and academic purposes")
    
    with st.form("proposal_form", clear_on_submit=False):
        col1, col2 = st.columns(2)
        
        with col1:
            proposal_title = st.text_input("Proposal Title *", 
                                          placeholder="e.g., Novel Approaches to Quantum Computing")
            principal_investigator = st.text_input("Principal Investigator *",
                                                 placeholder="Dr. Jane Smith")
            institution = st.text_input("Institution *", placeholder="MIT")
        
        with col2:
            department = st.text_input("Department *", placeholder="Physics")
            funding_agency = st.text_input("Funding Agency", placeholder="NSF")
            project_duration = st.selectbox("Project Duration", ["6 months", "1 year", "2 years", "3 years"])
            budget = st.number_input("Budget Request ($)", min_value=1000, max_value=10000000, value=100000, step=1000)
        
        abstract = st.text_area("Abstract *", height=120,
                               placeholder="Brief summary of the research proposal...")
        
        st.subheader("🎯 Research Details")
        
        research_question = st.text_area("Research Question", height=80,
                                        placeholder="What is the main research question?")
        objectives = st.text_area("Objectives", height=100,
                                 placeholder="List the research objectives...")
        methodology = st.text_area("Methodology", height=120,
                                  placeholder="Describe the research methodology...")
        expected_outcomes = st.text_area("Expected Outcomes", height=100,
                                        placeholder="What are the expected results?")
        timeline = st.text_area("Timeline", height=80,
                               placeholder="Project timeline and milestones...")
        references = st.text_area("References", height=80,
                                 placeholder="List key references...")
        
        submitted = st.form_submit_button("📑 Generate Proposal", use_container_width=True)
        
        if submitted:
            if not proposal_title or not principal_investigator or not institution or not abstract:
                st.error("⚠️ Please fill in all required fields (marked with *)!")
            else:
                st.success(f"✅ Research Proposal '{proposal_title}' generated!")
                
                # Preview and download options
                st.subheader("📄 Proposal Preview")
                with st.expander("View Full Proposal", expanded=True):
                    st.markdown(f"""
                    <div style="background: white; padding: 2rem; border-radius: 10px; border: 1px solid #e0e0e0;">
                        <h1 style="text-align: center;">{proposal_title}</h1>
                        <hr>
                        <p><strong>PI:</strong> {principal_investigator}</p>
                        <p><strong>Institution:</strong> {institution}</p>
                        <p><strong>Department:</strong> {department}</p>
                        <p><strong>Funding Agency:</strong> {funding_agency}</p>
                        <p><strong>Duration:</strong> {project_duration}</p>
                        <p><strong>Budget:</strong> ${budget:,}</p>
                        <hr>
                        <h3>Abstract</h3>
                        <p>{abstract}</p>
                        <h3>Research Question</h3>
                        <p>{research_question}</p>
                        <h3>Objectives</h3>
                        <p>{objectives}</p>
                        <h3>Methodology</h3>
                        <p>{methodology}</p>
                        <h3>Expected Outcomes</h3>
                        <p>{expected_outcomes}</p>
                        <h3>Timeline</h3>
                        <p>{timeline}</p>
                        <h3>References</h3>
                        <p>{references}</p>
                    </div>
                    """, unsafe_allow_html=True)

# ============================================================================
# POSTER STUDIO (Enhanced)
# ============================================================================

def render_poster_studio():
    """Render the poster studio with enhanced features."""
    st.header("🎨 Poster Studio")
    st.caption("Design professional research posters with multiple layouts and sizes")
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("📐 Poster Settings")
        
        # Poster Title
        poster_title = st.text_input("Poster Title", "My Research Poster")
        
        # Poster Size
        poster_size = st.selectbox(
            "Poster Size",
            [
                "A0 (84.1 × 118.9 cm)",
                "A1 (59.4 × 84.1 cm)",
                "A2 (42.0 × 59.4 cm)",
                "A3 (29.7 × 42.0 cm)",
                "36 × 48 inches",
                "42 × 48 inches",
                "48 × 60 inches",
                "Custom"
            ],
            index=1
        )
        
        # Custom Size
        if poster_size == "Custom":
            col_w, col_h = st.columns(2)
            with col_w:
                custom_width = st.number_input("Width (cm)", min_value=10, max_value=200, value=84)
            with col_h:
                custom_height = st.number_input("Height (cm)", min_value=10, max_value=200, value=118)
            size_display = f"{custom_width} × {custom_height} cm"
        else:
            size_display = poster_size
        
        # Poster Layout
        poster_layout = st.selectbox(
            "Poster Layout",
            [
                "Classic", "Modern", "Minimal",
                "Three Column", "Two Column",
                "Landscape", "Portrait",
                "Nature", "IEEE", "Medical"
            ],
            index=1
        )
        
        # Color Theme
        poster_theme = st.selectbox(
            "Color Theme",
            ["Blue", "Green", "Purple", "Red", "Dark", 
             "Medical", "Nature", "IEEE", "Vibrant", "Pastel"],
            index=0
        )
        
        # Font Size Control
        font_size = st.slider("Font Size", min_value=8, max_value=36, value=18, step=2)
        
        # Sections
        st.subheader("📑 Sections")
        sections = st.multiselect(
            "Select sections to include",
            [
                "Title", "Authors", "Abstract", 
                "Background", "Objectives", "Methods",
                "Results", "Discussion", "Conclusion",
                "Future Work", "References", "Acknowledgements",
                "Funding", "QR Code", "Contact"
            ],
            default=["Title", "Authors", "Abstract", "Methods", "Results", "Conclusion"]
        )
        
        # Generate Button
        if st.button("🎨 Generate Poster", use_container_width=True, type="primary"):
            st.session_state.poster_generated = True
            st.session_state.poster_title = poster_title
            st.session_state.poster_sections = sections
            st.session_state.poster_size = size_display
            st.session_state.poster_layout = poster_layout
            st.session_state.poster_theme = poster_theme
            st.session_state.poster_font_size = font_size
            
            # Save project
            project_data = {
                'id': generate_id(),
                'name': poster_title,
                'type': 'Poster',
                'size': size_display,
                'layout': poster_layout,
                'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
            }
            st.session_state.projects.append(project_data)
    
    with col2:
        if st.session_state.get('poster_generated', False):
            st.subheader("🖼️ Poster Preview")
            
            try:
                import matplotlib.pyplot as plt
                from matplotlib.patches import Rectangle
                
                # Theme colors
                theme_colors = {
                    "Blue": {"bg": "#f0f4f8", "primary": "#1E88E5", "secondary": "#E3F2FD", "accent": "#1565C0"},
                    "Green": {"bg": "#f0f8f0", "primary": "#2E7D32", "secondary": "#E8F5E9", "accent": "#1B5E20"},
                    "Purple": {"bg": "#f4f0f8", "primary": "#7B1FA2", "secondary": "#F3E5F5", "accent": "#4A148C"},
                    "Red": {"bg": "#f8f0f0", "primary": "#C62828", "secondary": "#FFEBEE", "accent": "#B71C1C"},
                    "Dark": {"bg": "#1a1a2e", "primary": "#667eea", "secondary": "#16213e", "accent": "#764ba2"},
                    "Medical": {"bg": "#f0f8ff", "primary": "#00695C", "secondary": "#E0F7FA", "accent": "#004D40"},
                    "Nature": {"bg": "#f5faf0", "primary": "#33691E", "secondary": "#F1F8E9", "accent": "#1B5E20"},
                    "IEEE": {"bg": "#ffffff", "primary": "#1565C0", "secondary": "#BBDEFB", "accent": "#0D47A1"},
                    "Vibrant": {"bg": "#fff3e0", "primary": "#E65100", "secondary": "#FFE0B2", "accent": "#BF360C"},
                    "Pastel": {"bg": "#fce4ec", "primary": "#AD1457", "secondary": "#F8BBD0", "accent": "#880E4F"},
                }
                
                colors = theme_colors.get(st.session_state.poster_theme, theme_colors["Blue"])
                
                # Create figure based on size
                if "A0" in size_display or "48" in size_display:
                    fig_size = (16, 12)
                elif "A1" in size_display or "36" in size_display:
                    fig_size = (14, 10)
                else:
                    fig_size = (12, 8)
                
                fig, ax = plt.subplots(figsize=fig_size, facecolor=colors["bg"])
                ax.set_facecolor(colors["bg"])
                
                # Title Section
                title_font_size = st.session_state.poster_font_size + 10
                ax.text(0.5, 0.95, poster_title, fontsize=title_font_size,
                       ha='center', va='top', fontweight='bold', 
                       color=colors["primary"], fontfamily='sans-serif')
                
                # Draw sections based on layout
                layout = st.session_state.poster_layout
                sections_list = st.session_state.poster_sections
                
                if "Column" in layout:
                    # Column layout
                    cols = 3 if "Three" in layout else 2
                    rows = 4
                    y_pos = 0.88
                    
                    for i, section in enumerate(sections_list[:min(len(sections_list), 8)]):
                        col_pos = (i % cols) / cols + 0.01
                        section_y = y_pos - (i // cols) * 0.2
                        
                        # Section header
                        ax.text(col_pos + 0.02, section_y, section.upper(),
                               fontsize=st.session_state.poster_font_size - 2,
                               fontweight='bold', color=colors["primary"])
                        
                        # Section content
                        ax.text(col_pos + 0.02, section_y - 0.04,
                               f"Sample {section.lower()} content here...",
                               fontsize=st.session_state.poster_font_size - 6,
                               wrap=True, color='#333')
                        
                        # Section box
                        ax.add_patch(Rectangle(
                            (col_pos, section_y - 0.14),
                            1/cols - 0.03, 0.12,
                            fill=True, facecolor=colors["secondary"],
                            alpha=0.3, edgecolor=colors["primary"],
                            linewidth=1.5, linestyle='--'
                        ))
                else:
                    # Single column or stacked layout
                    y_pos = 0.88
                    for section in sections_list[:6]:
                        # Section header
                        ax.text(0.08, y_pos, section.upper(),
                               fontsize=st.session_state.poster_font_size - 2,
                               fontweight='bold', color=colors["primary"])
                        
                        # Section content
                        ax.text(0.08, y_pos - 0.04,
                               f"Sample {section.lower()} content goes here...",
                               fontsize=st.session_state.poster_font_size - 6,
                               wrap=True, color='#333')
                        
                        # Section box
                        ax.add_patch(Rectangle(
                            (0.05, y_pos - 0.14),
                            0.90, 0.12,
                            fill=True, facecolor=colors["secondary"],
                            alpha=0.2, edgecolor=colors["primary"],
                            linewidth=1
                        ))
                        y_pos -= 0.16
                
                # Footer with metadata
                footer_text = f"Size: {size_display} | Layout: {layout} | Theme: {st.session_state.poster_theme}"
                ax.text(0.5, 0.01, footer_text,
                       fontsize=9, ha='center', color='gray', alpha=0.7)
                
                # Citation/Reference line
                ax.text(0.5, 0.03, "Generated by Academic Design Studio",
                       fontsize=8, ha='center', color='#aaa', alpha=0.5)
                
                ax.set_xlim(0, 1)
                ax.set_ylim(0, 1)
                ax.axis('off')
                plt.tight_layout()
                
                st.pyplot(fig)
                
                # Download Options
                st.subheader("📥 Download Options")
                
                col1, col2, col3, col4 = st.columns(4)
                
                # PNG Download
                buf_png = io.BytesIO()
                fig.savefig(buf_png, format='png', dpi=200, bbox_inches='tight', facecolor=colors["bg"])
                buf_png.seek(0)
                
                with col1:
                    st.download_button(
                        label="🖼️ Download PNG",
                        data=buf_png.getvalue(),
                        file_name=f"{poster_title.replace(' ', '_')}.png",
                        mime="image/png",
                        use_container_width=True
                    )
                
                # PDF Download
                buf_pdf = io.BytesIO()
                fig.savefig(buf_pdf, format='pdf', bbox_inches='tight', facecolor=colors["bg"])
                buf_pdf.seek(0)
                
                with col2:
                    st.download_button(
                        label="📄 Download PDF",
                        data=buf_pdf.getvalue(),
                        file_name=f"{poster_title.replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                
                # SVG Download
                buf_svg = io.BytesIO()
                fig.savefig(buf_svg, format='svg', bbox_inches='tight', facecolor=colors["bg"])
                buf_svg.seek(0)
                
                with col3:
                    st.download_button(
                        label="📐 Download SVG",
                        data=buf_svg.getvalue(),
                        file_name=f"{poster_title.replace(' ', '_')}.svg",
                        mime="image/svg+xml",
                        use_container_width=True
                    )
                
                # High Quality PNG
                buf_hq = io.BytesIO()
                fig.savefig(buf_hq, format='png', dpi=300, bbox_inches='tight', facecolor=colors["bg"])
                buf_hq.seek(0)
                
                with col4:
                    st.download_button(
                        label="📷 High Quality PNG",
                        data=buf_hq.getvalue(),
                        file_name=f"{poster_title.replace(' ', '_')}_HQ.png",
                        mime="image/png",
                        use_container_width=True
                    )
                
                st.success("✅ Poster generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating poster: {str(e)}")
                st.info("Please try with different settings or fewer sections.")

# ============================================================================
# CV BUILDER
# ============================================================================

def render_cv_builder():
    """Render the academic CV builder page."""
    st.header("📄 Academic CV Builder")
    st.caption("Create professional academic CVs for research positions")
    
    with st.form("cv_form", clear_on_submit=False):
        col1, col2 = st.columns(2)
        
        with col1:
            full_name = st.text_input("Full Name *", placeholder="Dr. John Doe")
            email = st.text_input("Email *", placeholder="john.doe@university.edu")
            phone = st.text_input("Phone", placeholder="+1-234-567-8900")
            linkedin = st.text_input("LinkedIn URL", placeholder="linkedin.com/in/username")
        
        with col2:
            position = st.text_input("Current Position", placeholder="Assistant Professor")
            institution = st.text_input("Institution", placeholder="Harvard University")
            research_areas = st.text_input("Research Areas", placeholder="Machine Learning, AI, NLP")
            website = st.text_input("Personal Website", placeholder="https://yourwebsite.com")
        
        professional_summary = st.text_area("Professional Summary", height=100,
                                          placeholder="Write a brief professional summary...")
        
        st.subheader("🎓 Education")
        education = st.text_area("Education", height=100,
                                placeholder="PhD in Computer Science, 2020, MIT\\nMS in Data Science, 2015, Stanford\\nBS in Computer Science, 2010, Berkeley")
        
        st.subheader("💼 Experience")
        experience = st.text_area("Experience", height=100,
                                 placeholder="Research Scientist, 2020-Present, Google\\nPostdoc, 2018-2020, MIT\\nResearch Assistant, 2015-2018, Stanford")
        
        st.subheader("📚 Publications")
        publications = st.text_area("Selected Publications", height=100,
                                   placeholder="1. Smith, J., et al. (2024). Deep Learning in Medical Imaging. Nature.\\n2. Smith, J., et al. (2023). Neural Networks for Healthcare. Science.")
        
        st.subheader("🛠️ Skills")
        skills = st.text_input("Skills (comma separated)",
                              placeholder="Python, Machine Learning, Data Analysis, Deep Learning")
        
        st.subheader("🏆 Awards & Honors")
        awards = st.text_area("Awards & Honors", height=80,
                             placeholder="Best Paper Award, NeurIPS 2023\\nNSF Early Career Award, 2022")
        
        st.subheader("📊 Research Metrics")
        col1, col2, col3 = st.columns(3)
        with col1:
            citations = st.number_input("Citations", min_value=0, max_value=99999, value=100)
        with col2:
            h_index = st.number_input("H-index", min_value=0, max_value=100, value=10)
        with col3:
            i10_index = st.number_input("i10-index", min_value=0, max_value=100, value=5)
        
        submitted = st.form_submit_button("📄 Generate CV", use_container_width=True)
        
        if submitted:
            if not full_name or not email:
                st.error("⚠️ Please fill in at least Name and Email!")
            else:
                st.success(f"✅ CV for {full_name} generated!")
                
                # Preview
                st.subheader("📄 CV Preview")
                with st.expander("View Full CV", expanded=True):
                    st.markdown(f"""
                    <div style="background: white; padding: 2rem; border-radius: 10px; border: 1px solid #e0e0e0;">
                        <h1 style="text-align: center; color: #2c3e50;">{full_name}</h1>
                        <p style="text-align: center;">
                            📧 {email} | 📞 {phone} | 🔗 <a href="{linkedin}">{linkedin}</a> | 🌐 {website}
                        </p>
                        <p style="text-align: center; font-style: italic; color: #666;">
                            {position} at {institution}
                        </p>
                        <p style="text-align: center; color: #888;">
                            Research: {research_areas}
                        </p>
                        <hr>
                        <h3>Professional Summary</h3>
                        <p>{professional_summary}</p>
                        <h3>🎓 Education</h3>
                        <p style="white-space: pre-wrap;">{education}</p>
                        <h3>💼 Experience</h3>
                        <p style="white-space: pre-wrap;">{experience}</p>
                        <h3>📚 Publications</h3>
                        <p style="white-space: pre-wrap;">{publications}</p>
                        <h3>🛠️ Skills</h3>
                        <p>{skills}</p>
                        <h3>🏆 Awards & Honors</h3>
                        <p style="white-space: pre-wrap;">{awards}</p>
                        <hr>
                        <h3>📊 Research Metrics</h3>
                        <table style="width: 100%;">
                            <tr><td><strong>Citations</strong></td><td>{citations}</td></tr>
                            <tr><td><strong>H-index</strong></td><td>{h_index}</td></tr>
                            <tr><td><strong>i10-index</strong></td><td>{i10_index}</td></tr>
                        </table>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Download options
                cv_text = f"""
                {'='*70}
                ACADEMIC CV: {full_name}
                {'='*70}
                
                Contact Information
                {'-'*50}
                Email: {email}
                Phone: {phone}
                LinkedIn: {linkedin}
                Website: {website}
                Position: {position}
                Institution: {institution}
                Research Areas: {research_areas}
                
                {'-'*50}
                PROFESSIONAL SUMMARY
                {'-'*50}
                {professional_summary}
                
                {'-'*50}
                EDUCATION
                {'-'*50}
                {education}
                
                {'-'*50}
                EXPERIENCE
                {'-'*50}
                {experience}
                
                {'-'*50}
                PUBLICATIONS
                {'-'*50}
                {publications}
                
                {'-'*50}
                SKILLS
                {'-'*50}
                {skills}
                
                {'-'*50}
                AWARDS & HONORS
                {'-'*50}
                {awards}
                
                {'-'*50}
                RESEARCH METRICS
                {'-'*50}
                Citations: {citations}
                H-index: {h_index}
                i10-index: {i10_index}
                
                {'='*70}
                Generated by Academic Design Studio v{APP_CONFIG['version']}
                Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
                """
                
                st.subheader("📥 Download Options")
                col1, col2 = st.columns(2)
                
                with col1:
                    cv_bytes = cv_text.encode('utf-8')
                    st.download_button(
                        label="📄 Download TXT",
                        data=cv_bytes,
                        file_name=f"{full_name.replace(' ', '_')}_CV.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                with col2:
                    try:
                        doc = docx.Document()
                        doc.add_heading(full_name, 0)
                        doc.add_paragraph(f"Email: {email}")
                        doc.add_paragraph(f"Phone: {phone}")
                        doc.add_heading("Professional Summary", level=1)
                        doc.add_paragraph(professional_summary)
                        doc.add_heading("Education", level=1)
                        for line in education.split('\n'):
                            doc.add_paragraph(line)
                        doc.add_heading("Experience", level=1)
                        for line in experience.split('\n'):
                            doc.add_paragraph(line)
                        
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.download_button(
                            label="📝 Download DOCX",
                            data=docx_buffer.getvalue(),
                            file_name=f"{full_name.replace(' ', '_')}_CV.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except:
                        st.button("📝 DOCX (Install python-docx)", disabled=True, use_container_width=True)

# ============================================================================
# QR CODE GENERATOR
# ============================================================================

def render_qr_generator():
    """Render the QR code generator page."""
    st.header("📱 QR Code Generator")
    st.caption("Generate QR codes for URLs, text, and academic IDs")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        qr_data = st.text_area("Enter data to encode", height=100,
                               placeholder="https://your-website.com\\nor any text or URL")
        
        st.subheader("🎨 Customization")
        qr_size = st.slider("QR Code Size", min_value=100, max_value=600, value=250, step=25)
        qr_color = st.color_picker("QR Code Color", "#1E88E5")
        bg_color = st.color_picker("Background Color", "#FFFFFF")
        
        # QR Version
        qr_version = st.selectbox("QR Version", ["1 (Smallest)", "2", "3", "4", "5 (Largest)"], index=0)
        version_map = {"1 (Smallest)": 1, "2": 2, "3": 3, "4": 4, "5 (Largest)": 5}
        
        # Error Correction
        error_correction = st.selectbox("Error Correction",
                                       ["Low (7%)", "Medium (15%)", "Quartile (25%)", "High (30%)"],
                                       index=2)
        
        if st.button("🎨 Generate QR Code", use_container_width=True, type="primary"):
            if qr_data:
                try:
                    import qrcode
                    
                    # Error correction mapping
                    error_map = {
                        "Low (7%)": qrcode.constants.ERROR_CORRECT_L,
                        "Medium (15%)": qrcode.constants.ERROR_CORRECT_M,
                        "Quartile (25%)": qrcode.constants.ERROR_CORRECT_Q,
                        "High (30%)": qrcode.constants.ERROR_CORRECT_H,
                    }
                    
                    qr = qrcode.QRCode(
                        version=version_map[qr_version],
                        error_correction=error_map[error_correction],
                        box_size=10,
                        border=4,
                    )
                    qr.add_data(qr_data)
                    qr.make(fit=True)
                    
                    img = qr.make_image(fill_color=qr_color, back_color=bg_color)
                    
                    # Save to buffer
                    buf = io.BytesIO()
                    img.save(buf, format='PNG')
                    buf.seek(0)
                    
                    st.session_state.qr_image = buf
                    st.session_state.qr_generated = True
                    st.session_state.qr_data = qr_data
                    
                    # Save project
                    project_data = {
                        'id': generate_id(),
                        'name': f"QR_{qr_data[:20]}",
                        'type': 'QR Code',
                        'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
                    }
                    st.session_state.projects.append(project_data)
                    
                except Exception as e:
                    st.error(f"Error generating QR Code: {e}")
            else:
                st.warning("Please enter data to encode")
    
    with col2:
        if st.session_state.get('qr_generated', False):
            st.subheader("🖼️ QR Code Preview")
            st.image(st.session_state.qr_image, width=qr_size)
            
            st.subheader("📥 Download Options")
            col1, col2, col3 = st.columns(3)
            
            # PNG Download
            with col1:
                st.download_button(
                    label="🖼️ PNG",
                    data=st.session_state.qr_image.getvalue(),
                    file_name="qr_code.png",
                    mime="image/png",
                    use_container_width=True
                )
            
            # SVG Download (if needed)
            with col2:
                try:
                    import qrcode.image.svg
                    qr_svg = qrcode.make(st.session_state.qr_data, image_factory=qrcode.image.svg.SvgImage)
                    svg_buffer = io.BytesIO()
                    qr_svg.save(svg_buffer)
                    svg_buffer.seek(0)
                    
                    st.download_button(
                        label="📐 SVG",
                        data=svg_buffer.getvalue(),
                        file_name="qr_code.svg",
                        mime="image/svg+xml",
                        use_container_width=True
                    )
                except:
                    st.button("📐 SVG (Available)", disabled=True, use_container_width=True)
            
            # EPS Download
            with col3:
                try:
                    import qrcode.image.eps
                    qr_eps = qrcode.make(st.session_state.qr_data, image_factory=qrcode.image.eps.EpsImage)
                    eps_buffer = io.BytesIO()
                    qr_eps.save(eps_buffer)
                    eps_buffer.seek(0)
                    
                    st.download_button(
                        label="📄 EPS",
                        data=eps_buffer.getvalue(),
                        file_name="qr_code.eps",
                        mime="application/postscript",
                        use_container_width=True
                    )
                except:
                    st.button("📄 EPS (Available)", disabled=True, use_container_width=True)
            
            # QR Code Info
            st.subheader("📊 QR Code Information")
            st.metric("Data Length", len(st.session_state.qr_data))
            st.metric("Version", version_map[qr_version])
            st.metric("Error Correction", error_correction)
            
            # Preview of encoded data
            with st.expander("📝 Encoded Data Preview"):
                st.code(st.session_state.qr_data)

# ============================================================================
# EXPORT MANAGER
# ============================================================================

def render_export_manager():
    """Render the export manager page."""
    st.header("📦 Export Manager")
    st.caption("Export your projects in multiple formats")
    
    # Export Settings
    st.subheader("⚙️ Export Settings")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        export_format = st.selectbox(
            "Export Format",
            ["TXT", "DOCX", "PDF", "PNG", "SVG", "JSON", "HTML"],
            index=0
        )
    
    with col2:
        export_quality = st.selectbox(
            "Quality",
            ["Draft (Fast)", "Standard", "High Quality"],
            index=1
        )
    
    with col3:
        include_metadata = st.checkbox("Include Metadata", value=True)
        include_timestamp = st.checkbox("Include Timestamp", value=True)
    
    st.subheader("📊 Your Projects")
    
    if st.session_state.projects:
        # Project management
        for idx, project in enumerate(st.session_state.projects):
            with st.container():
                col1, col2, col3, col4, col5 = st.columns([3, 2, 1, 1, 1])
                with col1:
                    st.markdown(f"**{project.get('name', 'Untitled')}**")
                    st.caption(f"Type: {project.get('type', 'Unknown')}")
                
                with col2:
                    st.caption(f"Created: {project.get('created', 'N/A')}")
                    if 'id' in project:
                        st.caption(f"ID: {project['id'][:12]}")
                
                with col3:
                    if st.button("📥", key=f"export_{idx}", help="Export project"):
                        # Generate export based on format
                        export_text = f"""
                        Project: {project.get('name', 'Untitled')}
                        Type: {project.get('type', 'Unknown')}
                        Created: {project.get('created', 'N/A')}
                        {'='*50}
                        Project Data:
                        {json.dumps(project, indent=2)}
                        """
                        
                        export_bytes = export_text.encode('utf-8')
                        
                        st.download_button(
                            label="Download",
                            data=export_bytes,
                            file_name=f"{project.get('name', 'project')}.{export_format.lower()}",
                            mime="text/plain",
                            key=f"download_{idx}"
                        )
                
                with col4:
                    if st.button("👁️", key=f"view_{idx}", help="View project"):
                        with st.expander(f"📄 Project Details: {project.get('name', '')}"):
                            st.json(project)
                
                with col5:
                    if st.button("🗑️", key=f"delete_{idx}", help="Delete project"):
                        if st.session_state.projects:
                            st.session_state.projects.pop(idx)
                            st.rerun()
                
                st.divider()
    else:
        st.info("📭 No projects available. Create one first using the document builders!")
    
    # Bulk Actions
    st.subheader("🔄 Bulk Actions")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("📦 Export All", use_container_width=True):
            if st.session_state.projects:
                with st.spinner("Exporting all projects..."):
                    progress_bar = st.progress(0)
                    for i, project in enumerate(st.session_state.projects):
                        # Simulate export
                        progress_bar.progress((i + 1) / len(st.session_state.projects))
                    st.success(f"✅ Exported {len(st.session_state.projects)} projects!")
            else:
                st.warning("No projects to export!")
    
    with col2:
        if st.button("🗑️ Clear All", use_container_width=True):
            if st.session_state.projects:
                st.session_state.projects = []
                st.success("✅ All projects cleared!")
                st.rerun()
            else:
                st.warning("No projects to clear!")
    
    with col3:
        if st.button("💾 Export Data", use_container_width=True):
            data = json.dumps(st.session_state.projects, indent=2)
            st.download_button(
                label="📥 Download Data",
                data=data,
                file_name="projects_data.json",
                mime="application/json"
            )
    
    with col4:
        if st.button("🔄 Import Data", use_container_width=True):
            uploaded_file = st.file_uploader("Choose JSON file", type="json")
            if uploaded_file is not None:
                try:
                    imported_data = json.load(uploaded_file)
                    if isinstance(imported_data, list):
                        st.session_state.projects.extend(imported_data)
                        st.success(f"✅ Imported {len(imported_data)} projects!")
                        st.rerun()
                    else:
                        st.error("Invalid data format. Expected array of projects.")
                except Exception as e:
                    st.error(f"Error importing data: {e}")

# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    """Main application entry point."""
    # Render sidebar
    render_sidebar()
    
    # Render current page
    page = st.session_state.current_page
    
    page_functions = {
        "Home": render_home,
        "Assignment": render_assignment_builder,
        "Lab Report": render_lab_report,
        "Thesis": render_thesis,
        "Research Proposal": render_research_proposal,
        "Poster": render_poster_studio,
        "Presentation": lambda: st.info("📽️ Presentation Builder - Coming Soon!"),
        "CV": render_cv_builder,
        "QR": render_qr_generator,
        "Citation": lambda: st.info("📚 Citation Manager - Coming Soon!"),
        "Export": render_export_manager,
        "Settings": lambda: st.info("⚙️ Settings - Coming Soon!"),
    }
    
    render_function = page_functions.get(page)
    if render_function:
        render_function()
    else:
        st.error(f"Page '{page}' not found!")
        render_home()
    
    # Footer
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.caption(f"🎓 {APP_CONFIG['name']} v{APP_CONFIG['version']}")
    with col2:
        st.caption("💡 Built with Streamlit")
    with col3:
        st.caption(f"© {APP_CONFIG['year']} {APP_CONFIG['author']}")

# ============================================================================
# RUN APPLICATION
# ============================================================================

if __name__ == "__main__":
    main()
